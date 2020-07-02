from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import copy
import re
doc = Document('resume.docx')
lists = ["Range", "adaptible", "quick learner", "essential"]  

for paragraph in doc.paragraphs:
    for target in lists:
        if target in paragraph.text:  

            currRuns = copy.copy(paragraph.runs)   
            paragraph.runs.clear()

            for run in currRuns:
                if target in run.text:
                    words = re.split('(\W)', run.text)  
                    for word in words:
                        if word == target:
                            newRun = paragraph.add_run(word)
                            newRun.font.highlight_color = WD_COLOR_INDEX.YELLOW
                        else:
                            newRun = paragraph.add_run(word)
                            newRun.font.highlight_color = None
                else: # 
                    paragraph.runs.append(run)


doc.save('resume1.docx')
